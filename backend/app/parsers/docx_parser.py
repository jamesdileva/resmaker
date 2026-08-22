"""DOCX parser built on python-docx."""

import re
from pathlib import Path

from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as DocxParagraph
from docx import Document

from app.parsers.base import BaseParser, ParsedDocument, Paragraph


def _numbering_info(paragraph: DocxParagraph) -> tuple[bool, int]:
    """Extract bullet membership and list level from numbering XML.

    Returns (has_numbering, ilvl). Real bullets carry a numPr element;
    the indent level (ilvl) is 0-based.
    """
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False, 0
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return False, 0
    ilvl_element = num_pr.find(qn("w:ilvl"))
    if ilvl_element is None:
        return True, 0
    try:
        return True, int(ilvl_element.get(qn("w:val")) or 0)
    except ValueError:
        return True, 0


def _indent_level(paragraph: DocxParagraph) -> int:
    """Fallback nesting level derived from left indentation."""
    left = paragraph.paragraph_format.left_indent
    if left is None:
        return 0
    inches = left.inches
    if inches <= 0.1:
        return 0
    # Each ~0.25" of indent equals one additional level.
    level = round(inches / 0.25)
    return max(level - 1, 0)


class DocxParser(BaseParser):
    """Extracts text, styles, headings, and bullets from DOCX files."""

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        document: DocumentObject = Document(str(path))

        paragraphs: list[Paragraph] = []
        for docx_paragraph in document.paragraphs:
            paragraphs.extend(self._parse_paragraph(docx_paragraph))
        paragraphs = [p for p in paragraphs if p.text.strip() or p.is_bullet]

        return ParsedDocument(
            filename=path.name,
            file_type="docx",
            paragraphs=paragraphs,
        )

    def supported_types(self) -> list[str]:
        return ["docx"]

    def _parse_paragraph(self, paragraph: DocxParagraph) -> list[Paragraph]:
        """Parse one DOCX paragraph, splitting on embedded line breaks."""
        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        has_numbering, list_level = _numbering_info(paragraph)

        is_heading = False
        heading_level = None
        if style_name.lower().startswith("heading"):
            parts = style_name.split()
            if len(parts) > 1 and parts[-1].isdigit():
                heading_level = int(parts[-1])
            else:
                heading_level = 1
            is_heading = True

        is_bullet = has_numbering or "list bullet" in style_name.lower()
        bullet_level = 0
        if is_bullet:
            numbered_style = re.match(
                r"list bullet\s+(\d+)", style_name.lower()
            )
            if numbered_style is not None:
                bullet_level = int(numbered_style.group(1))
            elif has_numbering:
                bullet_level = list_level + 1
            else:
                bullet_level = _indent_level(paragraph) + 1

        # Soft line breaks (<w:br>) surface as newlines inside .text;
        # each line becomes its own paragraph.
        lines = [line.strip() for line in paragraph.text.split("\n")]
        return [
            Paragraph(
                text=line,
                style=style_name,
                is_bullet=is_bullet,
                bullet_level=bullet_level,
                is_heading=is_heading,
                heading_level=heading_level,
            )
            for line in lines
            if line
        ]
