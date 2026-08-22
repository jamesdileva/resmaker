"""Document registry and export pipeline (DOCX/TXT)."""

import os
import tempfile
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt

from app.models.build import BuiltDocument


class DocumentRegistry:
    """In-memory store of built documents keyed by document_id.

    Documents are produced synchronously and returned to the client;
    the registry lets /export/ resolve them within the server process,
    mirroring the import job registry pattern.
    """

    def __init__(self) -> None:
        self._documents: dict[str, BuiltDocument] = {}

    def register(self, document: BuiltDocument) -> None:
        self._documents[document.document_id] = document

    def get(self, document_id: str) -> Optional[BuiltDocument]:
        return self._documents.get(document_id)


# Shared process-wide registry instance.
registry = DocumentRegistry()


class TxtExporter:
    """Plain-text output: UPPERCASE headers, '- ' bullets."""

    def export(self, document: BuiltDocument, include_traceability: bool = False) -> str:
        lines: list[str] = []
        for section in document.sections:
            lines.append(section.title.upper())
            lines.append("")
            for line in section.profile_lines:
                lines.append(line)
            for group in section.groups:
                header = group.title
                if group.dates:
                    header = f"{header} ({group.dates})"
                lines.append(header.upper())
                for bullet in group.bullets:
                    lines.append(f"- {bullet}")
            for line in section.lines:
                lines.append(f"- {line}")
            lines.append("")

        if include_traceability and document.traceability:
            lines.append("TRACEABILITY REPORT")
            lines.append("")
            for item_id, evidence_id in document.traceability.items():
                lines.append(f"- {item_id} -> {evidence_id}")

        return "\n".join(lines).strip() + "\n"


class DocxExporter:
    """DOCX output with headings, bullets, and optional traceability."""

    def __init__(self, font: str = "Calibri", font_size: int = 11) -> None:
        self.font = font
        self.font_size = font_size

    def export(self, document: BuiltDocument, include_traceability: bool = True) -> bytes:
        doc = DocxDocument()

        style = doc.styles["Normal"]
        style.font.name = self.font
        style.font.size = Pt(self.font_size)

        for section in document.sections:
            doc.add_heading(section.title, level=1)

            if section.profile_lines:
                first = doc.paragraphs[-1] if doc.paragraphs else None
                for line in section.profile_lines:
                    paragraph = doc.add_paragraph(line)
                    if first is not None and line == section.profile_lines[0]:
                        # Name line stands out slightly.
                        paragraph.runs[0].bold = True

            for group in section.groups:
                doc.add_heading(group.title, level=2)
                if group.dates:
                    doc.add_paragraph(group.dates)
                self._add_bullets(doc, group.bullets, document, include_traceability)

            if section.lines:
                self._add_bullets(doc, section.lines, document, include_traceability)

        if include_traceability and document.traceability:
            doc.add_heading("Traceability Report", level=1)
            for item_id, evidence_id in document.traceability.items():
                doc.add_paragraph(f"{item_id} -> {evidence_id}", style="List Bullet")

        return _save_to_bytes(doc)

    def _add_bullets(
        self,
        doc: DocxDocument,
        bullets: list[str],
        document: BuiltDocument,
        include_traceability: bool,
    ) -> None:
        evidence_by_item = {}
        for item_id, evidence_id in document.traceability.items():
            evidence_by_item[evidence_id] = evidence_id
        for bullet in bullets:
            paragraph = doc.add_paragraph(bullet, style="List Bullet")
            if include_traceability:
                evidence_id = self._find_evidence(bullet, document)
                if evidence_id:
                    self._attach_comment(
                        doc, paragraph, f"evidence:{evidence_id}"
                    )

    @staticmethod
    def _find_evidence(bullet_text: str, document: BuiltDocument) -> Optional[str]:
        # Bullets carry the content of the item they came from; find any
        # item whose id maps to evidence and whose content matches.
        for section in document.sections:
            for group in section.groups:
                if bullet_text in group.bullets:
                    return group.evidence_id
        return None

    @staticmethod
    def _attach_comment(doc: DocxDocument, paragraph, text: str) -> None:
        """Attach a Word comment; no-op if the runtime lacks the API."""
        try:
            add_comment = getattr(paragraph, "add_comment", None)
            if callable(add_comment):
                add_comment(text=text, author="Career OS", initials="COS")
        except Exception:
            # Traceability comments are best-effort metadata.
            pass


def _save_to_bytes(doc: DocxDocument) -> bytes:
    import io

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def save_exported(content: bytes | str, extension: str) -> dict:
    """Persist exported bytes to a temp file; return path and size."""
    suffix = f".{extension}"
    if isinstance(content, str):
        content_bytes = content.encode("utf-8")
    else:
        content_bytes = content
    handle = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    handle.write(content_bytes)
    handle.close()
    return {
        "file_path": handle.name,
        "file_size": os.path.getsize(handle.name),
    }
