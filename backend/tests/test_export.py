"""Tests for the export service and endpoint (Sprint 27)."""

import os

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlmodel import Session

from app.models.build import BuiltDocument
from app.models.resume import ExperienceGroup, RenderedSection
from app.services.export_service import (
    DocxExporter,
    TxtExporter,
    registry,
)


def _sample_document() -> BuiltDocument:
    return BuiltDocument(
        document_id="doc-test-1",
        template_name="standard",
        sections=[
            RenderedSection(
                title="Summary",
                section_type="profile",
                profile_lines=["John Doe", "john@example.com"],
            ),
            RenderedSection(
                title="Experience",
                section_type="experience",
                groups=[
                    ExperienceGroup(
                        evidence_id="ev-1",
                        title="Sales Associate",
                        dates="2019 - 2022",
                        bullets=["Resolved customer complaints daily"],
                    )
                ],
            ),
        ],
        traceability={"item-1": "ev-1"},
        warnings=[],
    )


# --- TxtExporter ---


def test_txt_export_uppercases_headers_and_uses_bullets() -> None:
    text = TxtExporter().export(_sample_document())
    lines = text.splitlines()
    assert "SUMMARY" in lines
    assert "SALES ASSOCIATE (2019 - 2022)" in lines
    assert "- Resolved customer complaints daily" in lines
    assert "John Doe" in text


def test_txt_export_traceability_optional() -> None:
    document = _sample_document()
    without_report = TxtExporter().export(document, include_traceability=False)
    with_report = TxtExporter().export(document, include_traceability=True)

    assert "TRACEABILITY REPORT" not in without_report
    assert "TRACEABILITY REPORT" in with_report
    assert "item-1 -> ev-1" in with_report


# --- DocxExporter ---


def test_docx_export_produces_valid_document_with_content() -> None:
    content = DocxExporter().export(_sample_document(), include_traceability=False)

    import io

    parsed = DocxDocument(io.BytesIO(content))
    headings = [p.text for p in parsed.paragraphs if p.style.name.startswith("Heading")]
    assert "Summary" in headings
    assert "Experience" in headings
    bullet_texts = [
        p.text for p in parsed.paragraphs if p.style.name == "List Bullet"
    ]
    assert "Resolved customer complaints daily" in bullet_texts


def test_docx_export_includes_traceability_section_when_requested() -> None:
    with_trace = DocxExporter().export(_sample_document(), include_traceability=True)
    without_trace = DocxExporter().export(_sample_document(), include_traceability=False)

    import io

    def all_text(content: bytes) -> str:
        return "\n".join(p.text for p in DocxDocument(io.BytesIO(content)).paragraphs)

    assert "Traceability Report" in all_text(with_trace)
    assert "item-1 -> ev-1" in all_text(with_trace)
    assert "Traceability Report" not in all_text(without_trace)


# --- Registry ---


def test_registry_round_trip() -> None:
    document = _sample_document()
    registry.register(document)
    assert registry.get("doc-test-1") is document
    assert registry.get("missing") is None


# --- Export API ---


def test_export_endpoint_docx(client) -> None:
    registry.register(_sample_document())

    response = client.post(
        "/api/v1/export/",
        json={
            "document_id": "doc-test-1",
            "format": "docx",
            "include_traceability": True,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["file_path"].endswith(".docx")
    assert body["file_size"] > 0
    assert os.path.exists(body["file_path"])

    os.remove(body["file_path"])


def test_export_endpoint_txt(client) -> None:
    registry.register(_sample_document())

    response = client.post(
        "/api/v1/export/",
        json={"document_id": "doc-test-1", "format": "txt"},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["file_path"].endswith(".txt")

    with open(body["file_path"], encoding="utf-8") as handle:
        assert "SUMMARY" in handle.read()

    os.remove(body["file_path"])


def test_export_endpoint_unknown_document_404(client) -> None:
    response = client.post(
        "/api/v1/export/", json={"document_id": "missing", "format": "txt"}
    )
    assert response.status_code == 404


def test_export_endpoint_rejects_pdf_for_now(client) -> None:
    response = client.post(
        "/api/v1/export/", json={"document_id": "whatever", "format": "pdf"}
    )
    # Unknown document check happens first; register one to hit format validation.
    registry.register(_sample_document())
    response = client.post(
        "/api/v1/export/", json={"document_id": "doc-test-1", "format": "pdf"}
    )
    assert response.status_code == 400


def test_built_resume_is_exportable_end_to_end(client, api_session: Session) -> None:
    """A resume built through /build/resume can be exported by document_id."""
    from app.db.models import Evidence, KnowledgeItem, KnowledgeItemEvidenceLink

    evidence = Evidence(
        type="experience",
        title="Boost Mobile",
        content="retail",
        role="Sales Associate",
        company="Boost Mobile",
    )
    item = KnowledgeItem(
        type="resume_bullet",
        content="Handled confidential records daily",
    )
    api_session.add(evidence)
    api_session.flush()
    api_session.add(item)
    api_session.flush()
    api_session.add(
        KnowledgeItemEvidenceLink(knowledge_item_id=item.id, evidence_id=evidence.id)
    )
    api_session.commit()

    built = client.post(
        "/api/v1/build/resume",
        json={"item_ids": [item.id], "user_profile": {"name": "John"}},
    ).json()
    assert built["document_id"]

    exported = client.post(
        "/api/v1/export/",
        json={"document_id": built["document_id"], "format": "docx"},
    )
    assert exported.status_code == 200
    assert os.path.exists(exported.json()["file_path"])
    os.remove(exported.json()["file_path"])
