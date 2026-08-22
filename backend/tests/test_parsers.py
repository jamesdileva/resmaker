"""Tests for the DOCX parser (Sprint 8)."""

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from app.parsers.base import BaseParser
from app.parsers.docx_parser import DocxParser
from app.parsers.pdf_parser import PdfParser
from app.parsers.txt_parser import TxtParser

FIXTURES = Path(__file__).resolve().parent / "fixtures"


def _add_numbering(paragraph, ilvl: str | None = None) -> None:
    """Attach a numPr element (optionally with an ilvl) to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.makeelement(qn("w:numPr"), {})
    if ilvl is not None:
        ilvl_element = p_pr.makeelement(qn("w:ilvl"), {qn("w:val"): ilvl})
        num_pr.append(ilvl_element)
    p_pr.append(num_pr)


def _write_doc(path: Path, builder) -> str:
    document = Document()
    builder(document)
    document.save(str(path))
    return str(path)


@pytest.fixture()
def resume_doc() -> str:
    return str(FIXTURES / "sample_resume.docx")


@pytest.fixture()
def soq_doc() -> str:
    return str(FIXTURES / "sample_soq.docx")


def test_parser_implements_base_interface() -> None:
    parser = DocxParser()
    assert isinstance(parser, BaseParser)
    assert parser.supported_types() == ["docx"]


def test_parse_returns_parsed_document(resume_doc: str) -> None:
    result = DocxParser().parse(resume_doc)
    assert result.filename == "sample_resume.docx"
    assert result.file_type == "docx"
    assert len(result.paragraphs) > 0


def test_headings_identified_with_levels(resume_doc: str) -> None:
    result = DocxParser().parse(resume_doc)
    headings = [p for p in result.paragraphs if p.is_heading]
    assert len(headings) == 3

    name_heading = headings[0]
    assert name_heading.text == "John Doe"
    assert name_heading.heading_level == 1

    job_headings = [h for h in headings if h.heading_level == 2]
    assert len(job_headings) == 2
    assert "Boost Mobile" in job_headings[0].text
    assert "County Office" in job_headings[1].text


def test_bullets_identified_from_styles(resume_doc: str) -> None:
    result = DocxParser().parse(resume_doc)
    bullets = [p for p in result.paragraphs if p.is_bullet]
    assert len(bullets) == 6
    assert all(not b.is_heading for b in bullets)
    assert all(b.bullet_level >= 1 for b in bullets)


def test_nested_bullet_level_from_style(resume_doc: str) -> None:
    result = DocxParser().parse(resume_doc)
    nested = [p for p in result.paragraphs if p.bullet_level >= 2]
    assert len(nested) == 1
    assert "cash drawers" in nested[0].text
    assert nested[0].style == "List Bullet 2"


def test_confidential_bullet_text_preserved(resume_doc: str) -> None:
    result = DocxParser().parse(resume_doc)
    matches = [
        p
        for p in result.paragraphs
        if "confidential customer records" in p.text
    ]
    assert len(matches) == 1
    assert matches[0].is_bullet


def test_soq_question_answer_structure(soq_doc: str) -> None:
    result = DocxParser().parse(soq_doc)
    questions = [
        p for p in result.paragraphs if p.is_heading and "Question" in p.text
    ]
    assert len(questions) == 2
    assert all(q.heading_level == 2 for q in questions)

    answers = [
        p
        for p in result.paragraphs
        if not p.is_heading and not p.is_bullet and p.text.strip()
    ]
    answer_texts = " ".join(a.text for a in answers)
    assert "confidential" in answer_texts
    assert "analysis of intake reports" in answer_texts


def test_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        DocxParser().parse(str(tmp_path / "does_not_exist.docx"))


# --- PDF parser (Sprint 9) ---


@pytest.fixture()
def pdf_doc() -> str:
    return str(FIXTURES / "sample_posting.pdf")


def test_pdf_parser_implements_interface(pdf_doc: str) -> None:
    parser = PdfParser()
    assert isinstance(parser, BaseParser)
    assert parser.supported_types() == ["pdf"]


def test_pdf_parse_extracts_all_text(pdf_doc: str) -> None:
    result = PdfParser().parse(pdf_doc)
    assert result.filename == "sample_posting.pdf"
    assert result.file_type == "pdf"
    assert len(result.paragraphs) >= 3
    joined = " ".join(p.text for p in result.paragraphs)
    assert "Office Technician" in joined
    assert "confidential files" in joined
    assert "written communication skills" in joined


def test_pdf_paragraphs_are_normal_style(pdf_doc: str) -> None:
    result = PdfParser().parse(pdf_doc)
    assert all(p.style == "Normal" for p in result.paragraphs)
    assert all(not p.is_bullet and not p.is_heading for p in result.paragraphs)


def test_missing_pdf_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        PdfParser().parse(str(tmp_path / "missing.pdf"))


# --- TXT parser (Sprint 9) ---


@pytest.fixture()
def duty_txt() -> str:
    return str(FIXTURES / "sample_duty.txt")


def test_txt_parser_implements_interface(duty_txt: str) -> None:
    parser = TxtParser()
    assert isinstance(parser, BaseParser)
    assert parser.supported_types() == ["txt"]


def test_txt_parse_returns_all_lines(duty_txt: str) -> None:
    result = TxtParser().parse(duty_txt)
    assert result.filename == "sample_duty.txt"
    assert result.file_type == "txt"

    texts = [p.text for p in result.paragraphs]
    # Blank lines are skipped; all content lines are present.
    assert len(texts) == 5
    assert texts[0] == "Duty Statement - Office Technician"
    assert any("confidential documents" in t for t in texts)


def test_txt_lines_are_normal_style(duty_txt: str) -> None:
    result = TxtParser().parse(duty_txt)
    assert all(p.style == "Normal" for p in result.paragraphs)
    assert all(not p.is_bullet and not p.is_heading for p in result.paragraphs)


def test_missing_txt_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        TxtParser().parse(str(tmp_path / "missing.txt"))


# --- Parser factory (Sprint 9) ---


def test_get_parser_dispatches_by_type() -> None:
    from app.parsers import get_parser

    assert isinstance(get_parser("docx"), DocxParser)
    assert isinstance(get_parser("pdf"), PdfParser)
    assert isinstance(get_parser("txt"), TxtParser)


def test_get_parser_is_case_insensitive_and_strips_dot() -> None:
    from app.parsers import get_parser

    assert isinstance(get_parser("DOCX"), DocxParser)
    assert isinstance(get_parser(".Pdf"), PdfParser)


def test_get_parser_unknown_raises_value_error() -> None:
    from app.parsers import get_parser

    with pytest.raises(ValueError):
        get_parser("xlsx")


def test_bullet_detected_from_numbering_xml(tmp_path: Path) -> None:
    def builder(document) -> None:
        paragraph = document.add_paragraph("Numbered but plain style")
        _add_numbering(paragraph, ilvl="1")

    parsed = DocxParser().parse(
        _write_doc(tmp_path / "numbered.docx", builder)
    )
    bullets = [p for p in parsed.paragraphs if p.is_bullet]
    assert len(bullets) == 1
    assert bullets[0].bullet_level == 2


def test_numbering_without_ilvl_defaults_to_level_one(tmp_path: Path) -> None:
    def builder(document) -> None:
        paragraph = document.add_paragraph("Bare numbering")
        _add_numbering(paragraph)

    parsed = DocxParser().parse(
        _write_doc(tmp_path / "bare.docx", builder)
    )
    bullets = [p for p in parsed.paragraphs if p.is_bullet]
    assert len(bullets) == 1
    assert bullets[0].bullet_level == 1


def test_malformed_ilvl_falls_back_to_level_one(tmp_path: Path) -> None:
    def builder(document) -> None:
        paragraph = document.add_paragraph("Broken level")
        _add_numbering(paragraph, ilvl="not-a-number")

    parsed = DocxParser().parse(
        _write_doc(tmp_path / "broken.docx", builder)
    )
    bullets = [p for p in parsed.paragraphs if p.is_bullet]
    assert len(bullets) == 1
    assert bullets[0].bullet_level == 1


def test_indentation_drives_bullet_level_without_paragraph_numbering(
    tmp_path: Path,
) -> None:
    def builder(document) -> None:
        # Style says bullet, but no paragraph-level numPr -> indentation decides.
        shallow = document.add_paragraph("Indented once", style="List Bullet")
        shallow.paragraph_format.left_indent = Inches(0.5)
        deep = document.add_paragraph("Indented deeper", style="List Bullet")
        deep.paragraph_format.left_indent = Inches(1.0)
        flush = document.add_paragraph("Flush bullet", style="List Bullet")
        flush.paragraph_format.left_indent = Inches(0.05)

    parsed = DocxParser().parse(
        _write_doc(tmp_path / "indents.docx", builder)
    )
    by_text = {p.text: p for p in parsed.paragraphs}
    assert by_text["Indented once"].is_bullet
    assert by_text["Indented once"].bullet_level == 2
    assert by_text["Indented deeper"].bullet_level == 4
    assert by_text["Flush bullet"].bullet_level == 1
