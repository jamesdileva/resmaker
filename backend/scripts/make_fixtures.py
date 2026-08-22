"""Generate DOCX/PDF/TXT test fixtures for the parser test suite."""

from pathlib import Path

import pymupdf
from docx import Document

FIXTURES_DIR = Path(__file__).resolve().parent.parent / "tests" / "fixtures"


def make_sample_resume(path: Path) -> None:
    document = Document()

    document.add_heading("John Doe", level=1)
    document.add_paragraph(
        "Customer service professional with 5+ years of experience."
    )

    document.add_heading("Sales Associate - Boost Mobile (2019-2022)", level=2)
    bullets_1 = [
        "Handled confidential customer records and account verification daily",
        "Resolved customer complaints, maintaining a 95% satisfaction rating",
        "Trained 4 new employees on point-of-sale and inventory systems",
    ]
    for text in bullets_1:
        document.add_paragraph(text, style="List Bullet")

    nested = document.add_paragraph(
        "Processed payments and reconciled cash drawers", style="List Bullet 2"
    )
    assert nested.style.name == "List Bullet 2"

    document.add_heading("Data Entry Clerk - County Office (2017-2019)", level=2)
    bullets_2 = [
        "Performed data analysis on weekly intake reports using Excel",
        "Maintained filing systems for over 500 sensitive documents",
    ]
    for text in bullets_2:
        document.add_paragraph(text, style="List Bullet")

    document.add_paragraph("Proficient in Excel, SQL dashboards, and CRM tools.")
    document.save(str(path))


def make_sample_soq(path: Path) -> None:
    document = Document()

    document.add_heading("Statement of Qualifications", level=1)
    document.add_paragraph("Applicant: John Doe")

    document.add_heading(
        "Question 1: Describe your experience handling confidential information",
        level=2,
    )
    document.add_paragraph(
        "Throughout my five years at Boost Mobile I managed confidential "
        "customer records, verifying identities and protecting private data "
        "in compliance with company privacy policies."
    )

    document.add_heading(
        "Question 2: Describe your analytical experience",
        level=2,
    )
    document.add_paragraph(
        "As a data entry clerk I performed weekly analysis of intake reports, "
        "built Excel dashboards, and presented summary findings to management."
    )

    document.save(str(path))


def make_sample_pdf(path: Path) -> None:
    document = pymupdf.open()
    page = document.new_page()

    blocks = [
        "Office Technician - State Department of Public Works",
        "",
        "Duties include reviewing confidential files, preparing weekly "
        "reports, and answering customer inquiries in person and by phone.",
        "",
        "Qualifications: two years of clerical experience, proficiency "
        "with spreadsheets, and strong written communication skills.",
    ]
    # Render each block into a wrapped textbox so nothing clips at the
    # right margin.
    rect = pymupdf.Rect(72, 72, 540, 400)
    page.insert_textbox(rect, "\n".join(blocks), fontsize=11, align=0)
    document.save(str(path))
    document.close()


def make_sample_duty_txt(path: Path) -> None:
    lines = [
        "Duty Statement - Office Technician",
        "",
        "1. Review and process confidential documents daily.",
        "2. Prepare weekly statistical reports using Excel.",
        "3. Answer customer inquiries via phone and email.",
        "4. Maintain office filing systems and supply inventory.",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    resume_path = FIXTURES_DIR / "sample_resume.docx"
    soq_path = FIXTURES_DIR / "sample_soq.docx"
    pdf_path = FIXTURES_DIR / "sample_posting.pdf"
    duty_path = FIXTURES_DIR / "sample_duty.txt"
    make_sample_resume(resume_path)
    make_sample_soq(soq_path)
    make_sample_pdf(pdf_path)
    make_sample_duty_txt(duty_path)
    print(f"Created {resume_path}")
    print(f"Created {soq_path}")
    print(f"Created {pdf_path}")
    print(f"Created {duty_path}")


if __name__ == "__main__":
    main()
